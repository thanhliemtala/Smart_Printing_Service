from datetime import datetime
from flask import Flask, render_template, flash, redirect, url_for, request,session
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.datastructures import  FileStorage
from form import RegistrationForm, LoginForm
from flask import jsonify
from flask_bcrypt import Bcrypt
from flask import Flask, request
from flask_login import LoginManager, UserMixin, login_user  
from flask_login import login_required
import os
import fitz
import docx
import txt

# Define current_stage globally
current_stage = 'print_document'
# Define stage_history as a global variable
stage_history = []
# Initialize uploaded_file globally
uploaded_file = None

app = Flask(__name__) #name is special variable
# Config app
app.config['SECRET_KEY'] = '80ea24d9322a81681369a178028405b8'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///app.db'
app.config["MAX_CONTENT_PATH"] = 100

# Create an instance database
db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
login_manager  = LoginManager(app)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

class User(db.Model,UserMixin):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    documents = db.relationship('Document', backref='user', lazy=True)
    
    def __init__(self, username, email, password):
        self.username = username
        self.email = email
        self.password = password
    
    def __repr__(self):
        return f"User({self.username},{self.email})"
    
class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    file_type = db.Column(db.String(10), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    date = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

    def __init__(self, filename, file_type,file_size, user_id):
        self.filename = filename
        self.file_type = file_type
        self.file_size = file_size
        self.user_id = user_id

    def __repr__(self):
        return f"Document({self.filename})"


class Printer(db.Model):
   
    brand_name = db.Column(db.String(20), nullable=False, primary_key=True)
    is_on = db.Column(db.Boolean, default=False)
    location = db.Column(db.String(20), nullable=False)
    notes = db.Column(db.String(20), nullable=False)
    print_model = db.Column(db.String(20), nullable=False)
    def __init__(self, brand_name, print_model, is_on, location, notes):
        self.brand_name = brand_name
        self.print_model = print_model
        self.location = location
        self.is_on = is_on
        self.notes = notes

    def __repr__(self):
        return f"Brand_name:({self.brand_name}), Model:({self.print_model})"

class PrintedDocument(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    document_name = db.Column(db.String(100), nullable=False)
    pages_printed = db.Column(db.Integer, nullable=False)
    date_printed = db.Column(db.DateTime, nullable=False, default=datetime.utcnow())
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

    def __repr__(self):
        return f"PrintedDocument('{self.document_name}', '{self.pages_printed}', '{self.date_printed}')"


@app.route('/', methods=["GET", "POST"])
def init():
    if request.method == "POST":
        user_type = request.form.get('user_type')
        # Storing userType in the session
        session['user_type'] = user_type

        # Redirect to appropriate page based on user type
        if user_type == 'user':
            return redirect(url_for('login'))  # Redirect to user login
        elif user_type == 'spso':
            return redirect(url_for('login'))  # Redirect to SPSO login
    return render_template("init.html")
@app.route('/home')
def home():
    if "user_type" in session:
        user_type = session['user_type']

        # Redirect to appropriate page based on user type
        if user_type == 'user':
            return render_template("user_dashboard.html") 
        elif user_type == 'spso':
            return render_template("spso_dashboard.html")
# Login
@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        with app.app_context():
            user = User.query.filter_by(email = form.email.data).first()
            if user and bcrypt.check_password_hash(user.password, form.password.data):
                login_user(user, remember=form.remember.data)    
                if "user_type" in session:
                    user_type = session.get('user_type')
                    if user_type == 'user':
                        return render_template("user_dashboard.html")  
                    elif user_type == 'spso':
                        return render_template("spso_dashboard.html")  
            else:
                flash('Login Unsuccessful. Please check email and password', 'danger')
    return render_template('login.html', form=form)


    
# Log out
@app.route('/logout')
def logout():
    session.pop("user_type",None)
    return redirect(url_for("init"))

## Print logic 

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
FREE_PAGES_LIMIT = 10

# Define the function to check if the file has an allowed extension

def is_valid_file(file_content):
    file_extension = os.path.splitext(file_content.filename)[1]
    valid_extensions = ('.pdf', '.txt', '.docx')
    return file_extension in valid_extensions


#Function to find number of pages
def get_number_of_pages(file_path):
    if file_path.lower().endswith('.pdf'):
        with fitz.open(file_path) as doc:
            return doc.page_count
    elif file_path.lower().endswith('.docx'):
        with docx.Document(file_path) as doc:
            return sum(1 for _ in doc.paragraphs)
    elif file_path.lower().endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            return len(file.readlines())
    else:
        return 0

def is_file_pages_valid(file_path):
    try:
        number_of_pages = get_number_of_pages(file_path)
        return isinstance(number_of_pages, int) and number_of_pages > 0
    except Exception as e:
        print(f"Error validating page count: {e}")
        return False


def get_pdf_or_txt_page_count(file_path):
    # Reusing existing logic from get_number_of_pages
    return get_number_of_pages(file_path)

def get_docx_page_count(file_path):
    # Reusing existing logic from get_number_of_pages
    return get_number_of_pages(file_path)

@app.route('/buy_more_pages_action', methods=['POST'])
@login_required
def buy_more_pages_action():
    flash("You have bought more pages!", 'success')
    return redirect(url_for('print_document'))

def is_valid_file(file_content):
    try:
        mime_type, _ = mimetypes.guess_type(None, file_content)
        return mime_type in {'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'text/plain'}
    except Exception as e:
        print(f"Error checking file type: {e}")
        return False


@app.route('/print_document', methods=['GET', 'POST'])
@login_required
def print_document():
    global current_stage, stage_history, uploaded_file, warning, file_extension

    # Reset warning and file_extension variables
    warning = False
    file_extension = ""

    if request.method == 'POST':
        if 'file' in request.files:
            uploaded_file = request.files['file']

            if uploaded_file:
                try:
                    # Read the file content
                    file_content = uploaded_file.read()

                    # Check if the file content is valid
                    if not file_content or not is_valid_file(file_content):
                        current_stage = 'file_not_allowed'
                        flash("Error: Invalid file content. Please choose a valid file.", 'danger')
                    else:
                        # Continue with your processing logic
                        number_of_pages, is_valid = get_number_of_pages(file_content)

                        print(f"Number of pages: {number_of_pages}")

                        if not is_valid:
                            current_stage = 'file_not_allowed'
                            flash(f"Error: Invalid file content. Please choose a valid file.", 'danger')
                        else:
                            # Store information about the printed document in the database
                            printed_document = PrintedDocument(
                                document_name=uploaded_file.filename,
                                pages_printed=number_of_pages,
                                content=file_content,  # Store file content in the database
                                user=current_user
                            )

                            db.session.add(printed_document)
                            db.session.commit()

                            # Check if the number of pages is within the free limit
                            if number_of_pages <= FREE_PAGES_LIMIT:
                                stage_history.append(current_stage)
                                current_stage = 'upload_success'
                            else:
                                # Check file type for specific warning messages
                                if file_extension == '.pdf':
                                    warning = True
                                    flash("Warning: You have exceeded the limit of free pages for PDF files.", 'warning')
                                elif file_extension == '.txt':
                                    warning = True
                                    flash("Warning: You have exceeded the limit of free pages for TXT files.", 'warning')
                                else:
                                    flash(f"Warning: You have exceeded the limit of free pages for {file_extension} files.", 'warning')

                            # Redirect to the appropriate stage
                            return redirect(url_for(current_stage))
                except Exception as e:
                    current_stage = 'file_not_allowed'
                    flash(f"Error: An unexpected error occurred while processing the file. Details: {str(e)}", 'danger')
                    print(f"Unexpected error: {e}")
            else:
                current_stage = 'file_not_allowed'
                flash(f"Error: Invalid file.", 'danger')

    return render_template('print.html', current_stage=current_stage, show_buy_button=False, uploaded_file=uploaded_file, warning=warning, file_extension=file_extension)

# ...


@app.route('/printer_selection', methods=['GET', 'POST'])
def printer_selection():
    available_printers = []  # Replace with your logic to get available printers

    if not available_printers:
        # Redirect to the notification page if no printers are available
        flash("No printers available. Please connect a printer and try again.", 'danger')
        return render_template('no_printer_available.html')

    if request.method == 'POST':
        selected_printer = request.form.get('printer')

        if selected_printer in available_printers:
            # Move to the next stage
            global current_stage
            current_stage = 'printing'  # Update to 'printing' or your next stage
            return render_template('print.html', current_stage=current_stage, selected_printer=selected_printer)
        else:
            # Display a notification if the selected printer is not available
            flash("Selected printer is not available. Please choose a different printer.", 'danger')

    # Display available printers
    return render_template('printer_selection.html', available_printers=available_printers)
    

def determine_next_stage(current_stage):
    # Find the index of the current stage in the history
    current_stage_index = stage_history.index(current_stage) if current_stage in stage_history else -1

    if current_stage_index > 0:
        # If the current stage is not the first one in the history,
        # return the stage right behind it
        return stage_history[current_stage_index - 1]
    else:
        # If the current stage is the first one or not found in history,
        # return 'print_document' as the default next stage
        return 'print_document'

@app.route('/move_to_next_stage', methods=['GET'])
def move_to_next_stage():
    global current_stage, stage_history

    # Call the determine_next_stage function to get the next stage
    next_stage = determine_next_stage(current_stage)

    # Add the current stage to the history
    stage_history.append(current_stage)

    # Update the current stage to the next stage
    current_stage = next_stage

    # Redirect to the determined next stage
    return redirect(url_for(current_stage))


@app.route('/move_to_previous_stage', methods=['GET'])
def move_to_previous_stage():
    global current_stage, stage_history
    if len(stage_history) > 1:
        # Remove the current stage from history
        stage_history.pop()
        # Get the previous stage from history
        current_stage = stage_history[-1]
    else:
        # Handle the case where there is no previous stage
        current_stage = 'print_document'

    # Redirect to the previous or default stage
    return redirect(url_for(current_stage))


###spso
@app.route('/printer_management', methods = ['GET', 'POST'])
def printer_management():
    printers = Printer.query.all()

    return render_template('printer_management.html', printers=printers)
    
@app.route('/configuration', methods = ['GET', 'POST'])
def configuration():
    return render_template("configuration.html")

@app.route('/reports', methods = ['GET', 'POST'])
def reports():
    return render_template("reports.html")

###### 
@app.route('/update_status/<int:printer_id>', methods=['POST'])
def update_status(printer_id):
    printer = Printer.query.get_or_404(printer_id)
    new_status = request.json.get('is_on', False)
    printer.is_on = new_status
    db.session.commit()
    return jsonify({'message': 'Status updated successfully'})

@app.route('/update_notes/<int:printer_id>', methods=['POST'])
def update_notes(printer_id):
    printer = Printer.query.get_or_404(printer_id)
    new_notes = request.json.get('notes', '')
    printer.notes = new_notes
    db.session.commit()
    return jsonify({'message': 'Notes updated successfully'})



# Create a route to save the printer information
@app.route('/add_printer', methods=['POST'])
def save_printer():
    brand_name = request.form.get('brand_name')
    print_model = request.form.get('print_model')
    is_on = bool(request.form.get('is_on'))
    location = request.form.get('location')
    notes = request.form.get('notes')

    # Create a Printer instance
    printer = Printer(brand_name=brand_name, print_model=print_model, is_on=is_on, location=location, notes=notes)

    # Add the Printer instance to the session
    db.session.add(printer)

    # Commit the changes to the database
    db.session.commit()

    return render_template("configuration.html")



######
with app.app_context():
    db.create_all()
if __name__ == '__main__':
    app.run(debug=True)
